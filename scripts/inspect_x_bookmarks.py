import zipfile, re, json

# Build map of X-hash anchors to bibliography entries by inspecting the per-chapter docx
# The X-hash bookmarks are placed on bibliography entries by pandoc-crossref

def build_anchor_to_citekey_map(docx_path, bib_path):
    """Builds a map from X-hash bookmark names to citekeys by matching bib entry text."""
    import docx as dx
    
    doc = dx.Document(docx_path)
    
    # 1. Collect all X-hash bookmarks and their paragraph text
    bookmark_text = {}
    for p in doc.paragraphs:
        bm_names = []
        for child in p._p.iter():
            if child.tag.endswith('bookmarkStart'):
                name = child.get(dx.oxml.ns.qn('w:name'), '')
                if name.startswith('X') and len(name) > 10:
                    bm_names.append(name)
        if bm_names:
            text = p.text.strip()
            for bm in bm_names:
                bookmark_text[bm] = text
    
    # 2. Collect all hyperlink anchors that point to X-hash -> their display text [N]
    anchor_display = {}
    for p in doc.paragraphs:
        for child in p._p:
            if child.tag.endswith('hyperlink'):
                anchor = child.get(dx.oxml.ns.qn('w:anchor'), '')
                if anchor.startswith('X') and len(anchor) > 10:
                    link_text = "".join([n.text for n in child.iter() if n.tag.endswith('t') and n.text]).strip()
                    if link_text and anchor not in anchor_display:
                        anchor_display[anchor] = link_text
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for child in p._p:
                        if child.tag.endswith('hyperlink'):
                            anchor = child.get(dx.oxml.ns.qn('w:anchor'), '')
                            if anchor.startswith('X') and len(anchor) > 10:
                                link_text = "".join([n.text for n in child.iter() if n.tag.endswith('t') and n.text]).strip()
                                if link_text and anchor not in anchor_display:
                                    anchor_display[anchor] = link_text
    
    # 3. Load bib database for matching
    with open(bib_path, 'r', encoding='utf-8') as f:
        bib_text = f.read()
    
    entries = re.findall(r'@\w+\{(\w+),', bib_text)
    
    # 4. For anchors that DON'T have bookmark definitions, we need to figure out citekey
    # by looking at display text [N] and matching with Daftar Pustaka order
    
    # First, get the bibliography order from the document
    in_bib = False
    bib_entries_ordered = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.lower() in ['daftar pustaka', 'references', 'bibliography']:
            in_bib = True
            continue
        if in_bib and p.style and p.style.name in ['Heading 1', 'Heading 2']:
            in_bib = False
        if in_bib and text:
            # Extract number if present
            num_match = re.match(r'^\[?(\d+)\]?\s*(.*)', text)
            if num_match:
                bib_entries_ordered.append((int(num_match.group(1)), text))
            else:
                bib_entries_ordered.append((len(bib_entries_ordered) + 1, text))
    
    print(f"Found {len(bookmark_text)} X-hash bookmarks with text")
    print(f"Found {len(anchor_display)} X-hash hyperlink anchors")
    print(f"Found {len(bib_entries_ordered)} Daftar Pustaka entries")
    
    for anchor, display in anchor_display.items():
        num_match = re.match(r'\[(\d+)\]', display)
        if num_match:
            num = int(num_match.group(1))
            print(f"  Anchor {anchor[:20]}... displays as [{num}]")
            if num <= len(entries):
                print(f"    -> citekey: {entries[num-1]}")

build_anchor_to_citekey_map('CNCTA/output/TA-final-raw.docx', 'CNCTA/referensi.bib')
