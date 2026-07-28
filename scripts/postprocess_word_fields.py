import os
import re
import random
import string
import json
import docx
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def load_bib_database(bib_path):
    """Parses BibTeX file into a CSL-JSON itemData dictionary keyed by citekey."""
    if not os.path.exists(bib_path):
        return {}
    with open(bib_path, 'r', encoding='utf-8') as f:
        bib_text = f.read()
        
    bib_db = {}
    raw_entries = re.split(r'\n@', '\n' + bib_text)
    for raw in raw_entries:
        if not raw.strip():
            continue
        match = re.match(r'^(\w+)\s*\{\s*([^,\s]+)\s*,\s*(.*)', raw, re.DOTALL)
        if match:
            entry_type, citekey, body = match.groups()
            fields = {}
            for line in body.split('\n'):
                line = line.strip()
                field_match = re.match(r'^(\w+)\s*=\s*[\"\{](.*)[\"\}],?$', line)
                if field_match:
                    fk, fv = field_match.groups()
                    fields[fk.lower()] = fv.rstrip('},"').strip()
            
            csl_type = 'article-journal' if entry_type == 'article' else entry_type
            item_data = {
                'id': citekey,
                'type': csl_type,
                'title': fields.get('title', ''),
            }
            if 'author' in fields:
                authors = []
                for a in fields['author'].split(' and '):
                    parts = a.split(',')
                    if len(parts) == 2:
                        authors.append({'family': parts[0].strip(), 'given': parts[1].strip()})
                    else:
                        authors.append({'family': a.strip()})
                item_data['author'] = authors
                
            if 'year' in fields:
                item_data['issued'] = {'date-parts': [[fields['year']]]}
                
            if 'doi' in fields:
                item_data['DOI'] = fields['doi']
                
            bib_db[citekey] = item_data
    return bib_db

def create_zotero_field_elements(citekey, display_text, bib_db):
    """Creates XML runs for ADDIN ZOTERO_ITEM CSL_CITATION Field."""
    citation_id = ''.join(random.choices(string.ascii_letters + string.digits, k=8))
    item_data = bib_db.get(citekey, {
        "id": citekey,
        "type": "book",
        "title": citekey
    })
    
    csl_dict = {
        "citationID": citation_id,
        "properties": {
            "unsorted": False,
            "formattedCitation": display_text,
            "plainCitation": display_text,
            "noteIndex": 0
        },
        "citationItems": [
            {
                "id": citekey,
                "uris": [f"http://zotero.org/users/local/project/items/{citekey}"],
                "itemData": item_data
            }
        ],
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json"
    }
    
    instr_text = f" ADDIN ZOTERO_ITEM CSL_CITATION {json.dumps(csl_dict)} "
    
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)
    
    r2 = OxmlElement('w:r')
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = instr_text
    r2.append(it)
    
    r3 = OxmlElement('w:r')
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3.append(fc2)
    
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = display_text
    r4.append(t)
    
    r5 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5.append(fc3)
    
    return [r1, r2, r3, r4, r5]

def add_complex_field(paragraph, instr_text, default_value="1"):
    """Adds a complex field (w:fldChar) to paragraph."""
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)
    paragraph._p.append(r1)
    
    r2 = OxmlElement('w:r')
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = f" {instr_text} "
    r2.append(it)
    paragraph._p.append(r2)
    
    r3 = OxmlElement('w:r')
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3.append(fc2)
    paragraph._p.append(r3)
    
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = default_value
    r4.append(t)
    paragraph._p.append(r4)
    
    r5 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5.append(fc3)
    paragraph._p.append(r5)

def replace_with_complex_ref(parent_elem, child_elem, bookmark_name, display_text):
    """Replaces child_elem in parent_elem with a Native Word Complex Field for REF."""
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)
    
    r2 = OxmlElement('w:r')
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = f" REF {bookmark_name} \\h "
    r2.append(it)
    
    r3 = OxmlElement('w:r')
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3.append(fc2)
    
    r4 = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r4.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = display_text
    r4.append(t)
    
    r5 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5.append(fc3)
    
    idx = list(parent_elem).index(child_elem)
    parent_elem.remove(child_elem)
    parent_elem.insert(idx, r5)
    parent_elem.insert(idx, r4)
    parent_elem.insert(idx, r3)
    parent_elem.insert(idx, r2)
    parent_elem.insert(idx, r1)

def replace_with_zotero_field(parent_elem, child_elem, citekey, display_text, bib_db):
    """Replaces child_elem in parent_elem with ADDIN ZOTERO_ITEM CSL_CITATION Field."""
    elements = create_zotero_field_elements(citekey, display_text, bib_db)
    idx = list(parent_elem).index(child_elem)
    parent_elem.remove(child_elem)
    for elem in reversed(elements):
        parent_elem.insert(idx, elem)

def format_table(table, header_bg="C00000"):
    """Formats table with exact reference styling."""
    tblPr = table._tbl.tblPr
    table.style = 'Table Grid'
    
    jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>')
    tblPr.append(jc)
    
    borders_xml = r'''
    <w:tblBorders %s>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
        <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
        <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>
    ''' % nsdecls('w')
    
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
            
    tblPr.append(parse_xml(borders_xml))
    
    for i, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        
        if i == 0:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
            
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{header_bg}"/>')
                tcPr.append(shd)
                
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for p in cell.paragraphs:
                    p.style = "Tabel"
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
        else:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for p in cell.paragraphs:
                    p.style = "Tabel"

def ensure_sumber_style(doc):
    """Creates a custom 'Sumber' paragraph style if it does not exist (Font 11pt, Special: None)."""
    existing_styles = [s.name for s in doc.styles]
    if "Sumber" not in existing_styles:
        sumber_style = doc.styles.add_style("Sumber", docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        if "Normal" in existing_styles:
            sumber_style.base_style = doc.styles["Normal"]
        sumber_style.font.name = "Times New Roman"
        sumber_style.font.size = docx.shared.Pt(11)
        sumber_style.font.italic = True
        sumber_style.paragraph_format.space_before = docx.shared.Pt(2)
        sumber_style.paragraph_format.space_after = docx.shared.Pt(6)
        sumber_style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        sumber_style.paragraph_format.first_line_indent = docx.shared.Pt(0)
        sumber_style.paragraph_format.left_indent = docx.shared.Pt(0)
    else:
        sumber_style = doc.styles["Sumber"]
        sumber_style.font.size = docx.shared.Pt(11)
        sumber_style.font.italic = True
        sumber_style.paragraph_format.first_line_indent = docx.shared.Pt(0)
        sumber_style.paragraph_format.left_indent = docx.shared.Pt(0)
    return sumber_style

def check_style_has_numPr(doc, style_name):
    """Checks if a paragraph style has automatic list numbering (w:numPr)."""
    if style_name in [s.name for s in doc.styles]:
        s = doc.styles[style_name]
        pPr = s._element.pPr
        if pPr is not None and pPr.numPr is not None:
            return True
    return False

def iter_all_paragraphs(doc):
    """Recursively yields all paragraphs from document body and table cells."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def build_xhash_citekey_map(doc, bib_db):
    """Builds a map from X-hash anchors to citekeys using the Daftar Pustaka ordering.
    
    pandoc-crossref sometimes generates hash-based anchors (X33ae9f...) instead of
    ref-citekey anchors. We resolve them by:
    1. Collecting the ordered Daftar Pustaka entries in the document
    2. Matching entry text against known bib_db entries by author/title similarity
    3. Mapping [N] display text -> Nth entry -> citekey
    """
    # 1. Collect ordered bibliography entries from the document
    in_bib = False
    bib_entries_ordered = []  # list of paragraph texts in order
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.lower() in ['daftar pustaka', 'references', 'bibliography']:
            in_bib = True
            continue
        if in_bib and p.style and p.style.name in ['Heading 1', 'Heading 2']:
            in_bib = False
        if in_bib and text:
            # Strip leading [N] number
            clean = re.sub(r'^\s*\[?\d+\]?\s*', '', text).strip()
            bib_entries_ordered.append(clean)
    
    # 2. Match each bib entry to a citekey by finding best text overlap
    num_to_citekey = {}
    for idx, entry_text in enumerate(bib_entries_ordered):
        best_key = None
        best_score = 0
        for ckey, item_data in bib_db.items():
            title = item_data.get('title', '').lower()
            authors = item_data.get('author', [])
            author_names = ' '.join([a.get('family', '') for a in authors]).lower()
            entry_lower = entry_text.lower()
            
            score = 0
            if title and title[:30] in entry_lower:
                score += 10
            if author_names:
                for aname in author_names.split():
                    if len(aname) > 2 and aname in entry_lower:
                        score += 3
            if score > best_score:
                best_score = score
                best_key = ckey
        if best_key and best_score >= 3:
            num_to_citekey[idx + 1] = best_key
    
    # 3. Collect all X-hash anchors and their [N] display text
    xhash_to_citekey = {}
    for p in iter_all_paragraphs(doc):
        for child in p._p:
            if child.tag.endswith('hyperlink'):
                anchor = child.get(qn('w:anchor'), '')
                if anchor.startswith('X') and len(anchor) > 10:
                    link_text = "".join([n.text for n in child.iter() if n.tag.endswith('t') and n.text]).strip()
                    num_match = re.match(r'\[(\d+)\]', link_text)
                    if num_match:
                        num = int(num_match.group(1))
                        if num in num_to_citekey and anchor not in xhash_to_citekey:
                            xhash_to_citekey[anchor] = num_to_citekey[num]
    
    return xhash_to_citekey

def process_document(docx_path, output_path, bib_path='CNCTA/referensi.bib', xhash_map=None):
    doc = docx.Document(docx_path)
    
    ensure_sumber_style(doc)
    bib_db = load_bib_database(bib_path)
    dp_has_numbering = check_style_has_numPr(doc, "Daftar Pustaka")
    
    # Build X-hash map if not provided externally
    if xhash_map is None:
        xhash_map = build_xhash_citekey_map(doc, bib_db)
    
    bookmark_map = {} # anchor -> bookmark_name
    bookmark_id_counter = 5000
    
    # 0. Normalize Raw Pandoc Bookmarks (eq:, fig:, tbl:) to _Ref_ format & reposition equation bookmarks
    for p in iter_all_paragraphs(doc):
        for child in list(p._p.iter()):
            if child.tag.endswith('bookmarkStart'):
                name = child.get(qn('w:name'), '')
                if name.startswith(('eq:', 'fig:', 'tbl:', 'ref-')):
                    target = name.replace('ref-', '', 1)
                    clean_name = "_Ref_" + re.sub(r'[^a-zA-Z0-9_]', '_', target)
                    child.set(qn('w:name'), clean_name)
                    bookmark_map[target] = clean_name
                    bookmark_map[name] = clean_name

                    # For equation bookmarks (eq:), move bookmark to wrap ONLY equation number (e.g. (1.1)) inside oMath
                    if target.startswith('eq:'):
                        b_id = child.get(qn('w:id'))
                        b_start = child
                        b_end = None
                        for elem in p._p.iter():
                            if elem.tag.endswith('bookmarkEnd') and elem.get(qn('w:id')) == b_id:
                                b_end = elem
                                break
                        omath = p._p.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
                        if omath is not None and b_end is not None:
                            omath_children = list(omath)
                            start_idx = None
                            end_idx = None
                            for i in range(len(omath_children) - 1, -1, -1):
                                ch = omath_children[i]
                                texts = [node.text for node in ch.iter() if node.tag.endswith('t') and node.text]
                                if ')' in texts and end_idx is None:
                                    end_idx = i
                                elif '(' in texts and end_idx is not None:
                                    start_idx = i
                                    break
                            if start_idx is not None and end_idx is not None and (end_idx - start_idx) <= 5:
                                b_start.getparent().remove(b_start)
                                b_end.getparent().remove(b_end)
                                omath.insert(start_idx, b_start)
                                omath.insert(end_idx + 2, b_end)

    # 1. Recursive Scan for Hyperlinks -> Zotero Fields (ref-citekey) & Word REF Fields (fig:, tbl:, eq:)
    for p in iter_all_paragraphs(doc):
        for child in list(p._p):
            if child.tag.endswith('hyperlink'):
                anchor = child.get(qn('w:anchor'))
                if anchor:
                    link_text = "".join([node.text for node in child.iter() if node.tag.endswith('t') and node.text]).strip()
                    
                    # Handle ref- prefix
                    if anchor.startswith('ref-'):
                        target_key = anchor.replace('ref-', '', 1)
                        if target_key.startswith(('fig:', 'tbl:', 'eq:')):
                            clean_name = "_Ref_" + re.sub(r'[^a-zA-Z0-9_]', '_', target_key)
                            bookmark_map[target_key] = clean_name
                            if not link_text:
                                link_text = "Rujukan"
                            replace_with_complex_ref(p._p, child, clean_name, link_text)
                        else:
                            # Academic citation -> Zotero Field!
                            if not link_text:
                                link_text = "[1]"
                            replace_with_zotero_field(p._p, child, target_key, link_text, bib_db)
                            
                    elif anchor.startswith(('tbl:', 'fig:', 'eq:')):
                        clean_name = "_Ref_" + re.sub(r'[^a-zA-Z0-9_]', '_', anchor)
                        bookmark_map[anchor] = clean_name
                        if not link_text:
                            link_text = "Rujukan"
                        replace_with_complex_ref(p._p, child, clean_name, link_text)
                    
                    # Handle X-hash anchors (pandoc-crossref generated hash citations)
                    elif anchor.startswith('X') and len(anchor) > 10:
                        if anchor in xhash_map:
                            citekey = xhash_map[anchor]
                            if not link_text:
                                link_text = "[1]"
                            replace_with_zotero_field(p._p, child, citekey, link_text, bib_db)

    # 2. Process Captions for Tables and Images
    caption_counts = {"Tabel": 0, "Gambar": 0}
    
    img_caption_indices = set()
    for i, p in enumerate(doc.paragraphs):
        has_drawing = any(n.tag.endswith(('drawing', 'pict')) for n in p._p.iter())
        if has_drawing and i + 1 < len(doc.paragraphs):
            for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip():
                    img_caption_indices.add(j)
                    break

    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else ""
        text = "".join([node.text for node in p._p.iter() if node.tag.endswith('t') and node.text]).strip()
        
        prefix = None
        if style_name == "TableCaption" or text.startswith("Tabel "):
            prefix = "Tabel"
        elif i in img_caption_indices or style_name == "ImageCaption" or text.startswith("Gambar ") or (style_name == "Caption" and not text.startswith("Tabel ")):
            prefix = "Gambar"
            
        if prefix and text:
            caption_counts[prefix] += 1
            
            anchor_found = None
            prefix_tag = "tbl:" if prefix == "Tabel" else "fig:"
            for anchor, b_name in list(bookmark_map.items()):
                if anchor.startswith(prefix_tag):
                    anchor_found = anchor
                    break
            
            if anchor_found:
                b_name = bookmark_map.pop(anchor_found)
            else:
                b_name = f"_Ref_{prefix_tag.replace(':', '_')}{caption_counts[prefix]}"
                
            bookmark_id_counter += 1
            b_id = str(bookmark_id_counter)
            
            sumber_text = ""
            sumber_match = re.search(r'\((Sumber|Source)\s*:.*?\)', text, re.IGNORECASE)
            if sumber_match:
                sumber_text = sumber_match.group(0)
                main_text = text.replace(sumber_text, "").strip()
            else:
                main_text = text
                
            pattern = rf'^{prefix}\s+[\d\.]+\s*:?\s*(.*)'
            match = re.match(pattern, main_text)
            if match:
                rest_of_caption = match.group(1)
            else:
                rest_of_caption = main_text.replace(f"{prefix} ", "", 1)
                
            rest_of_caption = re.sub(r'^\s*:\s*', '', rest_of_caption).strip()
            if rest_of_caption:
                rest_of_caption = " " + rest_of_caption
                
            pPr = p._p.pPr
            p._p.clear()
            if pPr is not None:
                p._p.append(pPr)
            
            p.style = "Caption"
            
            # --- BOOKMARK START ---
            b_start = OxmlElement('w:bookmarkStart')
            b_start.set(qn('w:id'), b_id)
            b_start.set(qn('w:name'), b_name)
            p._p.append(b_start)
            
            r1 = OxmlElement('w:r')
            t1 = OxmlElement('w:t')
            t1.text = f"{prefix} "
            t1.set(qn('xml:space'), 'preserve')
            r1.append(t1)
            p._p.append(r1)
            
            add_complex_field(p, r'STYLEREF 1 \s', default_value="1")
            
            r_dot = OxmlElement('w:r')
            t_dot = OxmlElement('w:t')
            t_dot.text = "."
            r_dot.append(t_dot)
            p._p.append(r_dot)
            
            add_complex_field(p, f'SEQ {prefix} \\* ARABIC \\s 1', default_value=str(caption_counts[prefix]))
            
            # --- BOOKMARK END ---
            b_end = OxmlElement('w:bookmarkEnd')
            b_end.set(qn('w:id'), b_id)
            p._p.append(b_end)
            
            r_rest = OxmlElement('w:r')
            t_rest = OxmlElement('w:t')
            t_rest.text = rest_of_caption
            t_rest.set(qn('xml:space'), 'preserve')
            r_rest.append(t_rest)
            p._p.append(r_rest)
            
            if sumber_text:
                p_sumber = p.insert_paragraph_before()
                p._p.addnext(p_sumber._p)
                p_sumber.style = "Sumber"
                p_sumber.text = sumber_text
                p_sumber.paragraph_format.space_before = docx.shared.Pt(2)
                p_sumber.paragraph_format.space_after = docx.shared.Pt(6)
                p_sumber.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                p_sumber.paragraph_format.first_line_indent = docx.shared.Pt(0)
                p_sumber.paragraph_format.left_indent = docx.shared.Pt(0)
                for r in p_sumber.runs:
                    r.font.italic = True
                    r.font.size = docx.shared.Pt(11)

    # 3. Apply Contextual Heading Paragraph Styles & Bibliography Formatting
    current_heading_context = "H2 Normal"
    in_bibliography = False
    
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        text = p.text.strip()
        
        if style_name == "Heading 1":
            if text.lower() in ["daftar pustaka", "references", "bibliography"]:
                in_bibliography = True
            else:
                in_bibliography = False
            current_heading_context = "H2 Normal"
            continue
        elif style_name == "Heading 2":
            in_bibliography = False
            current_heading_context = "H2 Normal"
            continue
        elif style_name == "Heading 3":
            in_bibliography = False
            current_heading_context = "H3 Normal"
            continue
        elif style_name == "Heading 4":
            in_bibliography = False
            current_heading_context = "H4 Normal"
            continue
            
        if in_bibliography:
            if text and style_name not in ["Heading 1", "Heading 2"]:
                p.style = "Daftar Pustaka"
                if dp_has_numbering:
                    clean_text = re.sub(r'^\s*\[\d+\]\s*', '', text)
                    if clean_text != text:
                        p.text = clean_text
            continue
            
        if style_name in ["Normal", "FirstParagraph", "Body Text", "Compact"] and text:
            p.style = current_heading_context

    # 4. Process Standalone Source Citations ((Sumber: ...)) -> Apply Style 'Sumber'
    for p in doc.paragraphs:
        text = p.text.strip()
        if (text.startswith("(Sumber:") or text.startswith("Sumber:") or text.startswith("(Source:")) and p.style.name != "Sumber":
            p.style = "Sumber"
            p.paragraph_format.space_before = docx.shared.Pt(2)
            p.paragraph_format.space_after = docx.shared.Pt(6)
            p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = docx.shared.Pt(0)
            p.paragraph_format.left_indent = docx.shared.Pt(0)

            for r in p.runs:
                r.font.italic = True
                r.font.size = docx.shared.Pt(11)

    # 5. Format Tables
    for table in doc.tables:
        format_table(table, header_bg="C00000")

    doc.save(output_path)
    print(f"Berhasil memproses dokumen lengkap & seluruh caption/hyperlinks/Zotero -> {output_path}")

if __name__ == '__main__':
    import sys
    if len(sys.argv) >= 3:
        process_document(sys.argv[1], sys.argv[2])
    else:
        process_document('CNCTA/output/TA-final-raw.docx', 'CNCTA/output/TA-final-complete.docx')
