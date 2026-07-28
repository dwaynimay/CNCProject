import re
import docx

doc = docx.Document('CNCTA/output/TA-final-complete.docx')

# 1. Find all Zotero field citekeys
zotero_keys = set()
zotero_count = 0
for p in doc.paragraphs:
    xml = p._p.xml
    if 'ADDIN ZOTERO_ITEM' in xml:
        zotero_count += 1
        ids = re.findall(r'"id":\s*"([^"]+)"', xml)
        zotero_keys.update(ids)

print(f"=== ZOTERO FIELD ANALYSIS ===")
print(f"Total Zotero field paragraphs: {zotero_count}")
print(f"Unique citekeys in Zotero fields ({len(zotero_keys)}):")
for k in sorted(zotero_keys):
    print(f"  {k}")

# 2. Find remaining hyperlinks (unconverted)
remaining = []
for i, p in enumerate(doc.paragraphs):
    for child in p._p:
        if child.tag.endswith('hyperlink'):
            anchor = child.get(docx.oxml.ns.qn('w:anchor'), '')
            rid = child.get(docx.oxml.ns.qn('r:id'), '')
            link_text = "".join([n.text for n in child.iter() if n.tag.endswith('t') and n.text])
            remaining.append((i, anchor, rid, link_text))

print(f"\n=== REMAINING UNCONVERTED HYPERLINKS ===")
print(f"Total remaining hyperlinks: {len(remaining)}")
for idx, anc, rid, txt in remaining:
    print(f"  P{idx:03d}: anchor={anc!r} rid={rid!r} text={txt!r}")

# 3. Also check table cells for hyperlinks
table_remaining = []
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for child in p._p:
                    if child.tag.endswith('hyperlink'):
                        anchor = child.get(docx.oxml.ns.qn('w:anchor'), '')
                        rid = child.get(docx.oxml.ns.qn('r:id'), '')
                        link_text = "".join([n.text for n in child.iter() if n.tag.endswith('t') and n.text])
                        table_remaining.append((ti, ri, ci, anchor, rid, link_text))

print(f"\n=== REMAINING HYPERLINKS IN TABLES ===")
print(f"Total hyperlinks in tables: {len(table_remaining)}")
for ti, ri, ci, anc, rid, txt in table_remaining:
    print(f"  Table{ti} Row{ri} Cell{ci}: anchor={anc!r} rid={rid!r} text={txt!r}")

# 4. Cross-check with markdown source citations
import glob
md_citations = set()
for path in sorted(glob.glob('CNCTA/TA-Rombak/0*.md')):
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    md_citations.update(re.findall(r'\[@([\w\-]+)\]', content))

# Filter to actual bibliography citations (not fig:, tbl:, eq:)
bib_citations = set(c for c in md_citations if not c.startswith(('fig:', 'tbl:', 'eq:')))

missing = bib_citations - zotero_keys
extra = zotero_keys - bib_citations

print(f"\n=== CITATION COVERAGE ANALYSIS ===")
print(f"Unique bib citations in Markdown: {len(bib_citations)}")
print(f"Unique citekeys in Zotero fields: {len(zotero_keys)}")
if missing:
    print(f"MISSING from Zotero fields ({len(missing)}):")
    for m in sorted(missing):
        print(f"  [MISSING] {m}")
else:
    print("[PASS] All markdown citations are covered by Zotero fields!")
if extra:
    print(f"Extra in Zotero (not in markdown): {sorted(extra)}")
