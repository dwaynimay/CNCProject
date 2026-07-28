import docx

doc = docx.Document('CNCTA/output/TA-final-raw.docx')
print('=== INSPECTING PARAGRAPHS AFTER IMAGES & TABLES ===')

for i, p in enumerate(doc.paragraphs):
    has_drawing = any(n.tag.endswith(('drawing', 'pict')) for n in p._p.iter())
    if has_drawing:
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i+1]
            print(f"P{i:03d} (Image) -> P{i+1:03d} [{next_p.style.name}]: '{next_p.text.strip()}'")

print('\n=== INSPECTING TABLE CAPTION PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith(('Table:', 'Tabel:', 'Tabel ')):
        print(f"P{i:03d} [{p.style.name}]: '{p.text.strip()}'")
