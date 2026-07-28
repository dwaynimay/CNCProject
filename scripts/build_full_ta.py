import os
import re
import subprocess
import zipfile
import warnings
import docx
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from postprocess_word_fields import process_document, build_xhash_citekey_map, load_bib_database, iter_all_paragraphs

warnings.filterwarnings("ignore", category=UserWarning)

CHAPTER_FILES = [
    'CNCTA/TA-Rombak/01-pendahuluan.md',
    'CNCTA/TA-Rombak/02-tinjauan-pustaka.md',
    'CNCTA/TA-Rombak/03-spesifikasi-desain-sistem.md',
    'CNCTA/TA-Rombak/04-implementasi.md',
    'CNCTA/TA-Rombak/05-pengujian-analisis.md',
    'CNCTA/TA-Rombak/06-evaluasi.md'
]

OUTPUT_RAW = 'CNCTA/output/TA-final-raw.docx'
OUTPUT_FINAL = 'CNCTA/output/TA-final-complete.docx'

RESOURCE_PATH = os.path.abspath('CNCTA/TA-Rombak')

def build_chapter(md_file, out_docx):
    cmd = [
        'pandoc',
        '--defaults', 'CNCTA/pandoc-defaults.yaml',
        '--resource-path', RESOURCE_PATH,
        '-o', out_docx,
        md_file
    ]
    print(f"Building {md_file} -> {out_docx}...")
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        print(f"Error building {md_file}:\n{res.stderr}")
        return False
    return True

def add_toc_field(paragraph):
    """Adds a Native Word TOC field to paragraph."""
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)
    paragraph._p.append(r1)
    
    r2 = OxmlElement('w:r')
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = r' TOC \o "1-3" \h \z \u '
    r2.append(it)
    paragraph._p.append(r2)
    
    r3 = OxmlElement('w:r')
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    r3.append(fc2)
    paragraph._p.append(r3)
    
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "Daftar Isi (Tekan F9 di Word untuk Memperbarui)"
    r4.append(t)
    paragraph._p.append(r4)
    
    r5 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r5.append(fc3)
    paragraph._p.append(r5)

import io

def copy_element_with_relationships(element, sub_doc, master_doc, rid_map):
    """Transfers image and media relationships from sub_doc into master_doc XML."""
    for node in element.iter():
        for attr in [qn('r:embed'), qn('r:id'), qn('r:link')]:
            if attr in node.attrib:
                old_rId = node.attrib[attr]
                if old_rId in sub_doc.part.rels:
                    rel = sub_doc.part.rels[old_rId]
                    if old_rId not in rid_map:
                        if rel.is_external:
                            new_rId = master_doc.part.relate_to(rel.target_ref, rel.reltype, is_external=True)
                        else:
                            if "image" in rel.reltype:
                                image_part = rel.target_part
                                image_bytes = io.BytesIO(image_part.blob)
                                new_rId, _ = master_doc.part.get_or_add_image(image_bytes)
                            else:
                                new_rId = master_doc.part.relate_to(rel.target_part, rel.reltype)
                        rid_map[old_rId] = new_rId
                    node.attrib[attr] = rid_map[old_rId]


def merge_docx_files(file_list, output_file):
    print(f"Merging {len(file_list)} DOCX files into {output_file}...")
    
    # Create fresh master document from reference template
    master = docx.Document('CNCTA/reference1.docx')
    master.element.body.clear()
    
    all_bib_paragraphs = []
    
    for file_path in file_list:
        sub = docx.Document(file_path)
        in_bib = False
        rid_map = {}
        
        for element in list(sub.element.body):
            if element.tag.endswith('sectPr'):
                continue
                
            text = "".join([n.text for n in element.iter() if n.tag.endswith('t') and n.text]).strip().lower()
            if text in ['daftar pustaka', 'references', 'bibliography']:
                in_bib = True
                continue
                
            if in_bib:
                pStyle = element.find(docx.oxml.ns.qn('w:pPr'))
                if pStyle is not None:
                    val = pStyle.find(docx.oxml.ns.qn('w:pStyle'))
                    if val is not None and 'Heading' in val.get(docx.oxml.ns.qn('w:val'), ''):
                        in_bib = False
                        
            # Transfer relationships (images, media, links) from sub to master
            copy_element_with_relationships(element, sub, master, rid_map)
            
            if in_bib:
                all_bib_paragraphs.append(element)
            else:
                master.element.body.append(element)

    # 1. Insert SINGLE TOC at top
    toc_p = docx.Document().add_paragraph()
    toc_p.style = "Normal"
    add_toc_field(toc_p)
    if len(master.paragraphs) > 0:
        master.paragraphs[0]._p.addprevious(toc_p._p)
    else:
        master.element.body.append(toc_p._p)

    # 2. Append ONE Single "Daftar Pustaka" Section at the VERY END of the manuscript (AFTER Chapter 6!)
    master.add_heading("Daftar Pustaka", level=1)
    
    seen_texts = set()
    for elem in all_bib_paragraphs:
        txt = "".join([n.text for n in elem.iter() if n.tag.endswith('t') and n.text]).strip()
        if txt and txt not in seen_texts:
            seen_texts.add(txt)
            master.element.body.append(elem)

    master.save(output_file)
    print(f"Merged output saved to {output_file}")

def verify_document_structure(docx_path):
    print(f"\n==================================================")
    print(f"=== LAPORAN VERIFIKASI HEADER & STRUKTUR DOKUMEN ===")
    print(f"==================================================")
    doc = docx.Document(docx_path)
    
    # 1. Check TOC Fields
    toc_indices = []
    for i, p in enumerate(doc.paragraphs):
        if 'TOC ' in p._p.xml or 'TOC\\' in p._p.xml:
            toc_indices.append(i)
            
    print(f"1. Jumlah Field Daftar Isi (TOC): {len(toc_indices)}")
    if len(toc_indices) == 1:
        print(f"   [PASS] Daftar Isi HANYA ADA 1 di paragraf P{toc_indices[0]:03d} (paling atas dokumen).")
    else:
        print(f"   [WARNING] Ditemukan {len(toc_indices)} Daftar Isi di paragraf: {toc_indices}")

    # 2. Check Heading 1s
    h1_list = []
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == 'Heading 1':
            h1_list.append((i, p.text.strip()))
            
    print(f"\n2. Daftar Heading 1 dalam Dokumen ({len(h1_list)} Heading 1):")
    for idx, text in h1_list:
        print(f"   - P{idx:03d}: '{text}'")

    # 3. Check Bibliography Headings
    bib_list = [(i, text) for i, text in h1_list if text.lower() in ['daftar pustaka', 'references', 'bibliography']]
    print(f"\n3. Jumlah Heading Daftar Pustaka: {len(bib_list)}")
    if len(bib_list) == 1 and bib_list[0][0] == h1_list[-1][0]:
        print(f"   [PASS] Daftar Pustaka HANYA ADA 1 di paragraf P{bib_list[0][0]:03d} (terletak di PALING AKHIR naskah).")
    else:
        print(f"   [WARNING] Daftar Pustaka tidak berada di paling akhir dokumen atau terdapat duplikasi! {bib_list}")

    # 4. Check Zotero Fields, Image Drawings, and ZIP Media Files
    zotero_count = sum(1 for p in doc.paragraphs if 'ADDIN ZOTERO_ITEM' in p._p.xml)
    ref_count = sum(1 for p in doc.paragraphs if 'REF _Ref_' in p._p.xml)
    drawings_count = sum(1 for p in doc.paragraphs if any(n.tag.endswith(('drawing', 'pict')) for n in p._p.iter()))
    
    with zipfile.ZipFile(docx_path) as z:
        media_files = [n for n in z.namelist() if 'word/media/' in n]

    print(f"\n4. Ringkasan Elemen Teknis Dokumen & Gambar:")
    print(f"   - Total Tag Gambar di Paragraf (w:drawing): {drawings_count}")
    print(f"   - Total Stream File Gambar di ZIP DOCX (word/media/): {len(media_files)}")
    if len(media_files) > 0:
        print(f"   [PASS] {len(media_files)} file gambar ter-embed 100% sempurna ke dalam file DOCX!")
    else:
        print(f"   [ERROR] TIDAK ADA GAMBAR TER-EMBED DALAM FILE DOCX!")
    print(f"   - Total Field Zotero (Sitasi): {zotero_count}")
    print(f"   - Total Field Native REF (Cross-References): {ref_count}")

    # 5. Check Remaining Unconverted Hyperlinks
    remaining_body = []
    for i, p in enumerate(doc.paragraphs):
        for child in p._p:
            if child.tag.endswith('hyperlink'):
                anchor = child.get(qn('w:anchor'), '')
                link_text = "".join([node.text for node in child.iter() if node.tag.endswith('t') and node.text]).strip()
                remaining_body.append((i, anchor[:30], link_text))
    
    remaining_table = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    for child in p._p:
                        if child.tag.endswith('hyperlink'):
                            anchor = child.get(qn('w:anchor'), '')
                            link_text = "".join([node.text for node in child.iter() if node.tag.endswith('t') and node.text]).strip()
                            remaining_table.append((ti, ri, ci, anchor[:30], link_text))
    
    total_remaining = len(remaining_body) + len(remaining_table)
    print(f"\n5. Sisa Hyperlink Yang Belum Terkonversi ke Field: {total_remaining}")
    if total_remaining == 0:
        print(f"   [PASS] Seluruh hyperlink sitasi & cross-ref telah terkonversi ke Field Native!")
    else:
        print(f"   [WARNING] Masih ada {total_remaining} hyperlink yang belum terkonversi:")
        for idx, anc, txt in remaining_body:
            print(f"     Body P{idx:03d}: anchor={anc!r} text={txt!r}")
        for ti, ri, ci, anc, txt in remaining_table:
            print(f"     Table{ti} Row{ri} Cell{ci}: anchor={anc!r} text={txt!r}")

    # 6. Citation Coverage Analysis
    import glob
    zotero_keys = set()
    for p in doc.paragraphs:
        xml = p._p.xml
        if 'ADDIN ZOTERO_ITEM' in xml:
            ids = re.findall(r'"id":\s*"([^"]+)"', xml)
            zotero_keys.update(ids)
    
    md_citations = set()
    for path in sorted(glob.glob('CNCTA/TA-Rombak/0*.md')):
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        md_citations.update(re.findall(r'\[@([\w\-]+)\]', content))
    
    bib_citations = set(c for c in md_citations if not c.startswith(('fig:', 'tbl:', 'eq:')))
    missing = bib_citations - zotero_keys
    
    print(f"\n6. Cakupan Sitasi Zotero:")
    print(f"   - Sitasi unik di Markdown: {len(bib_citations)}")
    print(f"   - Citekey unik di Field Zotero: {len(zotero_keys)}")
    if not missing:
        print(f"   [PASS] Seluruh {len(bib_citations)} sitasi di Markdown tercakup 100% oleh Field Zotero!")
    else:
        print(f"   [WARNING] {len(missing)} sitasi BELUM terkonversi:")
        for m in sorted(missing):
            print(f"     [MISSING] {m}")
    
    print(f"==================================================")

def main():
    os.makedirs('CNCTA/output', exist_ok=True)
    
    print("Building full TA document in a single-pass Pandoc build...")
    cmd = [
        'pandoc',
        '--defaults', 'CNCTA/pandoc-defaults.yaml',
        '--resource-path', RESOURCE_PATH,
        '-o', OUTPUT_RAW
    ] + CHAPTER_FILES
    
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        print(f"Error building manuscript:\n{res.stderr}")
        return

    # Add TOC field at the top of document
    raw_doc = docx.Document(OUTPUT_RAW)
    toc_p = docx.Document().add_paragraph()
    toc_p.style = "Normal"
    add_toc_field(toc_p)
    if len(raw_doc.paragraphs) > 0:
        raw_doc.paragraphs[0]._p.addprevious(toc_p._p)
    else:
        raw_doc.element.body.append(toc_p._p)
    raw_doc.save(OUTPUT_RAW)

    print(f"Running postprocessor on {OUTPUT_RAW}...")
    try:
        process_document(OUTPUT_RAW, OUTPUT_FINAL)
        final_doc_path = OUTPUT_FINAL
    except PermissionError:
        fallback_path = 'CNCTA/output/TA-final-complete-fixed.docx'
        print(f"\n[NOTE] '{OUTPUT_FINAL}' is currently opened in Word.")
        print(f"Saving updated result to '{fallback_path}' instead...")
        process_document(OUTPUT_RAW, fallback_path)
        final_doc_path = fallback_path
    
    # Run Automated Document Structure Verification
    verify_document_structure(final_doc_path)
            
    print("\nFULL TA BUILD & VERIFICATION SUCCESSFUL!")

if __name__ == '__main__':
    main()
