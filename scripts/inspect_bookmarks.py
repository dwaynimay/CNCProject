import docx

doc = docx.Document('CNCTA/output/TA-final-raw.docx')

# Find all bookmarkStart elements and their names
print("=== BOOKMARKS IN RAW DOCX ===")
for p in doc.paragraphs:
    for child in p._p.iter():
        if child.tag.endswith('bookmarkStart'):
            name = child.get(docx.oxml.ns.qn('w:name'), '')
            bid = child.get(docx.oxml.ns.qn('w:id'), '')
            if name.startswith('X') or name.startswith('ref-'):
                print(f"  Bookmark: name={name!r} id={bid} para_text={p.text[:60]!r}")

# Find Daftar Pustaka entries and their bookmarks
print("\n=== DAFTAR PUSTAKA ENTRIES WITH BOOKMARKS ===")
in_bib = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().lower() in ['daftar pustaka', 'references']:
        in_bib = True
        continue
    if in_bib and p.style and p.style.name in ['Heading 1', 'Heading 2']:
        in_bib = False
    if in_bib and p.text.strip():
        bookmarks = []
        for child in p._p.iter():
            if child.tag.endswith('bookmarkStart'):
                name = child.get(docx.oxml.ns.qn('w:name'), '')
                if name != '_GoBack':
                    bookmarks.append(name)
        if bookmarks:
            print(f"  P{i:03d}: bookmarks={bookmarks} text={p.text[:60]!r}")
